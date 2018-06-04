import docx
import re
from itertools import zip_longest
from config import cursor, db


def initDocx(name, path):
    path = '/root/exam_system/docx/' + path +'.docx'
    sql = 'insert into docx (paper_name, path) values (%s, %s)'
    cursor.execute(sql,(name, path))
    db.commit()
    makepaper()


def getDocxPath():
    sql = 'select * from docx order by id'
    cursor.execute(sql,)
    path = cursor.fetchone().get('path')
    return path

def getQuestion():
    doc = docx.Document(getDocxPath())

    paragraphs = doc.paragraphs
    temp = []
    for line in paragraphs:
        temp.append(line.text + '\n')
    exam = ''.join(temp)

    question = re.findall('[0-9]+、.*', exam)

    return question[0:46]


def getAnswer():
    doc = docx.Document(getDocxPath())
    paragraphs = doc.paragraphs
    temp = []
    for line in paragraphs:
        temp.append(line.text + '\n')
    exam = ''.join(temp)

    answer = re.findall('(【正确答案】.*|【答案】.*|参考答案：.*)', exam)[:46]
    answer_list = []
    for item in answer:
        if '【答案】' in item:
            answer_list.append(item.strip().replace(',', '')[4:])
        elif '【正确答案】' in item:
            answer_list.append(item.strip().replace(',', '')[6:])
        elif '参考答案：' in item:
            answer_list.append(item.strip().replace(',', '')[5:])

    return answer_list


def getItems():
    Items = []
    doc = docx.Document(getDocxPath())
    paragraphs = doc.paragraphs
    temp = []
    for line in paragraphs:
        temp.append(line.text + '\n')
    exam = ''.join(temp)
    items = re.findall('[A-D]\..*', exam)

    for i in range(0, len(items), 4):
        Items.append(items[i:i + 4])

    return Items[0:36]


def analyze():
    lists = []
    for question, items in zip_longest(getQuestion(), getItems()):
        lists.append({'question': question, 'items': items})
    return lists


def decidetra(a):
    if a == '×':
        return 0
    elif a == '√':
        return 1
    else:
        return a


def makepaper():
    for index, item in enumerate(analyze()):
        answer = list(map(decidetra, getAnswer()))[index]
        sql = 'select count(*) as last_id from docx'
        cursor.execute(sql)
        paper_id = cursor.fetchone().get('last_id')

        sql = 'insert into questions (q_text,q_type,q_value,A,B,C,D,paper_id,answer) values (%s,%s,%s,%s,%s,%s,%s,%s,%s)'

        if index <= 23:
            cursor.execute(sql, (
                item.get('question'), 'radio', 1.5, item.get('items')[0], item.get('items')[1], item.get('items')[2],
                item.get('items')[3], paper_id, answer))

            db.commit()

        if index > 23 and index < 36:
            cursor.execute(sql, (
                item.get('question'), 'checkbox', 2, item.get('items')[0], item.get('items')[1], item.get('items')[2],
                item.get('items')[3], paper_id, answer))

            db.commit()

        if index >= 36:
            cursor.execute(sql, (
                item.get('question'), 'decide', 1, '1', 0, 0,
                0, paper_id, answer))

            db.commit()


if __name__ == '__main__':
    initDocx('stefan', 'e:/test.docx')